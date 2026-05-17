"""Réécrit le template Word ``BilledE2E_parcours_administrateur.docx``.

Le template fournit la mise en page officielle (logo Billed, titre, tableau
bleu avec lignes Given/When/Then). Ce script remplace uniquement le contenu
textuel des 12 scénarios par ceux dérivés des tests Jest réellement écrits
sur le parcours administrateur (src/__tests__/Login.js partie admin,
src/__tests__/Dashboard.js, src/__tests__/Logout.js), tout en préservant
police, couleurs et structure du document d'origine.

Le script produit ensuite le PDF de livraison via LibreOffice headless.

Usage :
    python3 docs/generate_e2e_admin_plan.py
"""

from __future__ import annotations

import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_DOCX = REPO_ROOT / "docs" / "BilledE2E_parcours_administrateur_template.docx"
OUTPUT_DOCX = REPO_ROOT / "docs" / "BilledE2E_parcours_administrateur.docx"
OUTPUT_PDF = REPO_ROOT / "Henneron_Mathys_4_plan_test_admin_052026.pdf"


# 12 scénarios reflétant les tests Jest écrits pour le parcours admin.
# (given, when, then) — un tuple par scénario, dans l'ordre 1 → 12.
SCENARIOS: list[tuple[str, str, str]] = [
    # 1 — Login admin : champs vides
    # Test : Login.js "When I do not fill fields and I click on admin button"
    (
        "Je suis un visiteur (non connecté).",
        "Je ne remplis pas le champ e-mail ou le champ password du login "
        "administrateur et je clique sur le bouton \"Se connecter\".",
        "Je reste sur la page Login et je suis invité à remplir le champ "
        "manquant.",
    ),
    # 2 — Login admin : mauvais format d'e-mail
    # Test : Login.js "When I do fill fields in incorrect format and I click on admin button"
    (
        "Je suis un visiteur (non connecté).",
        "Je remplis le champ e-mail du login administrateur au mauvais "
        "format (sans la forme chaîne@chaîne) et je clique sur le bouton "
        "\"Se connecter\".",
        "Je reste sur la page Login et je suis invité à remplir le champ "
        "e-mail au bon format.",
    ),
    # 3 — Login admin : succès
    # Test : Login.js "When I do fill fields in correct format" + "It should renders HR dashboard page"
    (
        "Je suis un visiteur (non connecté).",
        "Je remplis le champ e-mail et le champ password du login "
        "administrateur au bon format et je clique sur le bouton "
        "\"Se connecter\" du formulaire Administration.",
        "Je suis authentifié en tant qu'administrateur RH (objet user de "
        "type Admin dans le localStorage) et je suis redirigé vers la "
        "page Dashboard (titre \"Validations\").",
    ),
    # 4 — Dashboard : affichage initial
    # Test : Dashboard.js intégration GET "fetches bills from mock API GET"
    (
        "Je suis connecté en tant qu'administrateur RH.",
        "J'arrive sur la page Dashboard.",
        "Je vois les sections \"En attente\", \"Validé\" et \"Refusé\" avec "
        "le compteur de notes de frais pour chaque statut, ainsi que la "
        "grande icône Billed à droite (zone \"Validations\").",
    ),
    # 5 — Dashboard : dépliage d'une section
    # Test : Dashboard.js "When I am on Dashboard page and I click on arrow"
    (
        "Je suis connecté en tant qu'administrateur sur la page Dashboard "
        "et au moins une note de frais existe pour le statut considéré.",
        "Je clique sur la flèche d'une section (\"En attente\", \"Validé\" "
        "ou \"Refusé\").",
        "La section se déplie et affiche la liste des cartes des notes de "
        "frais correspondantes (prénom/nom de l'employé, libellé, montant, "
        "date, type).",
    ),
    # 6 — Dashboard : ouverture du détail d'une note en attente
    # Test : Dashboard.js "When I am on Dashboard page and I click on edit icon of a card" → "right form should be filled"
    (
        "Je suis connecté en tant qu'administrateur et la section \"En "
        "attente\" est dépliée.",
        "Je clique sur la carte d'une note de frais en statut \"en "
        "attente\".",
        "Le formulaire de la note de frais s'affiche à droite avec "
        "l'ensemble des champs remplis (type, nom, date, commentaire, "
        "montant, TVA, justificatif) et les boutons \"Refuser\" et "
        "\"Accepter\" sont visibles.",
    ),
    # 7 — Dashboard : fermeture du détail (2e clic)
    # Test : Dashboard.js "When I am on Dashboard page and I click 2 times on edit icon" → "big bill Icon should Appear"
    (
        "Je suis connecté en tant qu'administrateur et le formulaire de "
        "détail d'une note de frais est ouvert.",
        "Je clique une deuxième fois sur la même carte de note de frais.",
        "Le formulaire de détail se ferme et la grande icône Billed "
        "(big-billed-icon) réapparaît à droite à la place du formulaire.",
    ),
    # 8 — Dashboard : validation (Accepter)
    # Test : Dashboard.js "When I click on accept button" → big-billed-icon
    (
        "Je suis connecté en tant qu'administrateur et le formulaire "
        "d'une note de frais au statut \"en attente\" est ouvert.",
        "Je saisis éventuellement un commentaire puis je clique sur le "
        "bouton \"Accepter\".",
        "Le statut de la note de frais passe à \"accepté\", le formulaire "
        "est remplacé par la grande icône Billed et le compteur \"Validé\" "
        "est incrémenté de 1 dans le feed des notes de frais.",
    ),
    # 9 — Dashboard : refus (Refuser)
    # Test : Dashboard.js "When I click on refuse button" → big-billed-icon
    (
        "Je suis connecté en tant qu'administrateur et le formulaire "
        "d'une note de frais au statut \"en attente\" est ouvert.",
        "Je saisis éventuellement un commentaire puis je clique sur le "
        "bouton \"Refuser\".",
        "Le statut de la note de frais passe à \"refusé\", le formulaire "
        "est remplacé par la grande icône Billed et le compteur \"Refusé\" "
        "est incrémenté de 1 dans le feed des notes de frais.",
    ),
    # 10 — Dashboard : visualisation du justificatif via la modale
    # Test : Dashboard.js "When I click on the icon eye" → modal opens
    (
        "Je suis connecté en tant qu'administrateur et le formulaire "
        "d'une note de frais est ouvert.",
        "Je clique sur l'icône en forme d'œil située dans le formulaire.",
        "Une modale s'ouvre et affiche le justificatif (image du reçu) "
        "associé à la note de frais ; la modale peut être fermée via le "
        "bouton \"×\".",
    ),
    # 11 — Dashboard : erreur API
    # Tests : Dashboard.js "fetches bills from an API and fails with 404 / 500 message error"
    (
        "Je suis connecté en tant qu'administrateur et l'API back-end "
        "renvoie une erreur (404 ou 500) sur l'appel GET /bills.",
        "Je navigue vers la page Dashboard.",
        "La page d'erreur est affichée avec le message correspondant "
        "(\"Erreur 404\" ou \"Erreur 500\") ; les sections de notes de "
        "frais ne sont pas chargées.",
    ),
    # 12 — Déconnexion
    # Test : Logout.js "When I click on disconnect button" → "I should be sent to login page"
    (
        "Je suis connecté en tant qu'administrateur sur la page Dashboard.",
        "Je clique sur l'icône de déconnexion (layout-disconnect) en bas "
        "de la barre verticale.",
        "Ma session est supprimée du localStorage et je suis redirigé "
        "vers la page Login (le formulaire \"Administration\" est de "
        "nouveau visible).",
    ),
]


def _set_cell_text_preserving_format(cell, text: str) -> None:
    """Remplace le texte d'une cellule sans toucher au style.

    Les cellules du template contiennent un seul paragraphe avec un (ou
    deux) runs. On écrit le nouveau texte dans le premier run pour
    conserver la police, la couleur, la taille, etc., et on vide les
    runs suivants.
    """

    paragraph = cell.paragraphs[0]
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def rewrite_template() -> Path:
    if not TEMPLATE_DOCX.exists():
        raise FileNotFoundError(
            f"Template introuvable : {TEMPLATE_DOCX}. Conservez le .docx "
            "fourni par OpenClassrooms sous ce chemin."
        )

    if len(SCENARIOS) != 12:
        raise ValueError(
            f"Le template attend exactement 12 scénarios, "
            f"{len(SCENARIOS)} fournis."
        )

    doc = Document(str(TEMPLATE_DOCX))
    table = doc.tables[0]

    for index, (given, when, then) in enumerate(SCENARIOS):
        base_row = index * 5  # 4 lignes par scénario + 1 ligne vide
        _set_cell_text_preserving_format(table.rows[base_row + 1].cells[1], given)
        _set_cell_text_preserving_format(table.rows[base_row + 2].cells[1], when)
        _set_cell_text_preserving_format(table.rows[base_row + 3].cells[1], then)

    doc.save(str(OUTPUT_DOCX))
    return OUTPUT_DOCX


def convert_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
    with tempfile.TemporaryDirectory() as tmpdir:
        profile = Path(tmpdir) / "lo_profile"
        outdir = Path(tmpdir) / "out"
        outdir.mkdir()
        subprocess.run(
            [
                "soffice",
                "--headless",
                f"-env:UserInstallation=file://{profile}",
                "--convert-to",
                "pdf",
                str(docx_path),
                "--outdir",
                str(outdir),
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.PIPE,
        )
        produced = outdir / (docx_path.stem + ".pdf")
        if not produced.exists():
            raise RuntimeError("LibreOffice n'a pas généré le PDF attendu.")
        shutil.copy2(produced, pdf_path)
    return pdf_path


def main() -> int:
    docx = rewrite_template()
    print(f"DOCX généré : {docx}")
    pdf = convert_to_pdf(docx, OUTPUT_PDF)
    print(f"PDF généré  : {pdf}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
